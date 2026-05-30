import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io
import fitz  # PyMuPDF
import copy

from docx.oxml.ns import qn

def find_element_and_copy(source_doc, marker_text, dest_doc, all_markers=None):
    """
    在 source_doc 中查找包含 marker_text 的段落，
    通常在标书中，模板会放在文档的最后面，而在前面的目录或要求说明中会提及。
    为了避免拷贝前面的说明文字，我们将寻找 marker_text 的【最后一次】有效出现位置，并从那里开始拷贝。
    """
    if not marker_text or marker_text.strip() in ['null', 'None']:
        return False
        
    clean_marker = marker_text.strip().replace(" ", "")
    if not clean_marker:
        return False
        
    other_markers = []
    if all_markers:
        other_markers = [m.strip().replace(" ", "") for m in all_markers if m and m.strip().replace(" ", "") != clean_marker]

    candidate_indices = []
    
    import re
    # 提取核心词汇（去掉括号及内部内容），增加容错率
    core_marker = re.sub(r'[（\(].*?[）\)]', '', clean_marker)
    if len(core_marker) < 4:
        core_marker = clean_marker[:8]
        
    # 首先遍历所有元素，找到所有可能的起始位置
    body_elements = list(source_doc.element.body)
    for idx, elem in enumerate(body_elements):
        if elem.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(elem, source_doc)
            clean_p = p.text.strip().replace(" ", "")
            if not clean_p: continue
            
            # 精确匹配或核心前缀匹配
            if clean_marker in clean_p or (len(core_marker) >= 4 and core_marker in clean_p):
                # 启发式判断：标题段落通常不会很长（排除掉大段的说明文字，但允许标题带有一长串括号说明）
                if len(clean_p) < len(clean_marker) + 60:
                    candidate_indices.append(idx)
                    
    # まず候補を評価する
    best_idx = -1
    best_score = -9999
    
    for c_idx in candidate_indices:
        score = 0
        elem = body_elements[c_idx]
        p = docx.text.paragraph.Paragraph(elem, source_doc)
        c_text = p.text.strip().replace(" ", "")
        
        # Exact match bonus
        if c_text == clean_marker or c_text == core_marker:
            score += 10
            
        # Penalize checklist items like "1.投标声明书......页码"
        if "...." in p.text or "页" in c_text[-3:] or c_text.startswith(clean_marker + "......"):
            score -= 50
            
        # Check the next 15 elements to see if it looks like a template body
        template_clues = 0
        has_table = False
        text_length = 0
        for next_i in range(c_idx + 1, min(c_idx + 16, len(body_elements))):
            next_elem = body_elements[next_i]
            if next_elem.tag.endswith('tbl'):
                has_table = True
            elif next_elem.tag.endswith('p'):
                next_p = docx.text.paragraph.Paragraph(next_elem, source_doc)
                nxt_txt = next_p.text.strip()
                text_length += len(nxt_txt)
                if any(w in nxt_txt for w in ['签字', '盖章', '日期', '年  月', '年   月', '法定代表人', '授权代表', '致：', '致:', '承诺', '___']):
                    template_clues += 1
                # If we hit another marker very quickly, it might be a checklist
                cln_nxt = nxt_txt.replace(" ", "")
                if cln_nxt and any((om in cln_nxt or cln_nxt in om) for om in other_markers):
                    if text_length < 50:  # Hit another marker with almost no content in between
                        score -= 20
                    break
        
        if has_table:
            score += 5
        score += template_clues * 3
        
        # Prefer candidates closer to the end of the document (appendices usually at the end)
        # c_idx / len(body_elements) is between 0 and 1
        position_bonus = (c_idx / len(body_elements)) * 5
        score += position_bonus
        
        if score > best_score:
            best_score = score
            best_idx = c_idx
            
    if best_idx == -1:
        return False
        
    start_idx = best_idx
    
    elements_copied = 0
    max_elements = 100 
    
    import re
    
    found_start = False
    
    sliding_window = []
    signature_block_detected = False
    lines_after_signature = 0
    
    for i in range(start_idx, len(body_elements)):
        elem = body_elements[i]
        # 仅处理段落和表格
        if not elem.tag.endswith(('p', 'tbl')):
            continue
            
        text = ""
        is_heading = False
        
        if elem.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(elem, source_doc)
            text = p.text.strip()
            is_heading = p.style.name.startswith('Heading')
            
        clean_text = text.replace(" ", "")
        
        just_found = False
        if not found_start:
            # 检查是否匹配模板标题
            if clean_marker in clean_text:
                found_start = True
                just_found = True
                
        if not found_start:
            continue
            
        # 判断是否是下一个模板的标题 (跳过刚刚匹配到的第一段)
        is_new_template = False
        if not just_found and text and elements_copied > 1:
            # 标题不应该以句号、分号、逗号、冒号结尾（排除模板正文中的正常段落或列表声明项）
            if not re.search(r'[。；;，,：:]$', text.strip()):
                if 2 < len(text) < 80:
                    # 1. 启发式正则
                    if re.match(r'^附件\s*\d+', text) or re.match(r'^附表\s*\d+', text):
                        is_new_template = True
                    elif text.startswith('格式') and len(text) < 15:
                        is_new_template = True
                    elif text.startswith('第') and '部分' in text:
                        is_new_template = True
                    # 响应用户的需求：遇到 (1) (2) 连续数字标号的独立段落，直接认为是新章节
                    elif re.match(r'^[（\(]?[一二三四五六七八九十\d]+[）\)、\.]', text):
                        # 核心防误杀：正文中的“1.我方保证提交的文件真实。”不应被当成新标题。将冒号也加入黑名单。
                        if not re.search(r'[，。；：:]', text):
                            # 如果没有特定关键词，则只要它不太长（比如典型的标题长度），就认为是新章节
                            if any(keyword in text for keyword in ['函', '表', '格式', '声明', '部分', '材料', '证明', '文件', '承诺', '报价', '书']):
                                is_new_template = True
                            elif len(text) < 25:
                                is_new_template = True
                    # 2. 绝对匹配其他已知模板的特征词 (维度1：下一个文件阻断)
                    if clean_text and any((om in clean_text or clean_text in om) for om in other_markers):
                        # 如果匹配了其他模板名字，但也需要确保它是个“标题”的样子，不能是一长串带标点的话
                        if not re.search(r'[，。；：:]', text):
                            is_new_template = True
                            
            if (is_heading and elements_copied > 3) or is_new_template:
                break

        # 深拷贝 XML 节点
        new_elem = copy.deepcopy(elem)
        
        # 剥离原有的分页符，避免克隆导致跳页 (使用 xpath 递归查找所有层级)
        for br in new_elem.xpath('.//w:br'):
            if br.get(qn('w:type')) == 'page':
                br.getparent().remove(br)
        for pbb in new_elem.xpath('.//w:pageBreakBefore'):
            pbb.getparent().remove(pbb)
            
        # 同时删除可能强制换页的 w:sectPr
        for sect in new_elem.xpath('.//w:sectPr'):
            sect.getparent().remove(sect)

        # 关键修复：不能直接 append，必须插入到 sectPr（如果有）之前，否则 Word 会把它们全部丢到文档最后或破坏结构
        sectPr = dest_doc.element.body.find(qn('w:sectPr'))
        if sectPr is not None:
            sectPr.addprevious(new_elem)
        else:
            dest_doc.element.body.append(new_elem)
            
        elements_copied += 1
        
        if elements_copied >= max_elements:
            break
            
    return elements_copied > 0

def create_styled_document(analysis_data: dict, volume_index: int, attachment_bytes: bytes = None, source_bytes: bytes = None, source_filename: str = None) -> io.BytesIO:
    """基于大模型提取的数据，生成原生带有样式排版的 Word 分册"""
    source_doc = None
    if source_bytes and source_filename and source_filename.lower().endswith('.docx'):
        try:
            # 极速克隆：直接使用原文档作为母本，清空内容，保留所有样式、边距、字体等，确保100%原汁原味
            doc = docx.Document(io.BytesIO(source_bytes))
            doc._element.body.clear_content()
            
            # 清空原文档中残留的页眉页脚（因为克隆会连带页眉页脚一起复制，包括首个页眉和偶数页眉）
            for sectPr in doc._element.body.findall(qn('w:sectPr')):
                for header_ref in sectPr.findall(qn('w:headerReference')):
                    sectPr.remove(header_ref)
                for footer_ref in sectPr.findall(qn('w:footerReference')):
                    sectPr.remove(footer_ref)
                        
            source_doc = docx.Document(io.BytesIO(source_bytes))
        except Exception as e:
            print(f"Error loading source docx: {e}")
            doc = docx.Document()
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.styles['Normal'].font.size = Pt(12)
    else:
        doc = docx.Document()
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)
    
    volumes = analysis_data.get('volumes', [])
    if volume_index < 0 or volume_index >= len(volumes):
        raise ValueError("Invalid volume index")
        
    vol = volumes[volume_index]
    
    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"\n\n\n\n\n【{vol.get('volume_name', '投标文件')}】\n\n")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    
    proj_name = analysis_data.get('projectInfo', {}).get('name', '未命名项目')
    proj_num = analysis_data.get('projectInfo', {}).get('number', '未知编号')
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"项目名称：{proj_name}\n项目编号：{proj_num}\n\n\n\n")
    sub_run.font.size = Pt(18)
    
    doc.add_page_break()
    
    # 提取所有的 marker 用于辅助防越界
    items = vol.get('items', [])
    all_markers = []
    for item in items:
        tt = item.get('template_text')
        if tt and tt.strip() not in ['null', 'None']:
            all_markers.append(tt)
            
    is_first_content = True
    for item in items:
        item_name = item.get('name', '未知材料')
        item_type = item.get('type', 'file')
        
        if item_type == 'heading':
            if not is_first_content:
                doc.add_page_break()
            # 这是一个分类标题，渲染为加粗的标题行
            h = doc.add_heading(item_name, level=2)
            is_first_content = False
            continue
            
        if not is_first_content:
            doc.add_page_break()
        doc.add_heading(item_name, level=3)
        is_first_content = False
        
        template_text = item.get('template_text')
        has_template = False
        
        if source_doc and template_text:
            # 尝试在源文档中查找并克隆
            has_template = find_element_and_copy(source_doc, template_text, doc, all_markers)
            
        if not has_template:
            # 留白或简单提示
            if template_text and template_text.strip() not in ['null', 'None']:
                p_info = doc.add_paragraph(f"[原文档模板特征词：{template_text}，未自动匹配到格式，请自行插入]")
                p_info.runs[0].font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            else:
                p_info = doc.add_paragraph("\n[请在此处插入对应的正文或扫描件]\n")
                p_info.runs[0].font.color.rgb = docx.shared.RGBColor(128, 128, 128)

    # 附加证明材料 (PDF 转图片)
    if attachment_bytes:
        doc.add_heading('附件：资质证明与附件扫描件', level=1)
        try:
            with fitz.open(stream=attachment_bytes, filetype="pdf") as pdf_doc:
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc.load_page(page_num)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("png")
                    
                    img_stream = io.BytesIO(img_bytes)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(img_stream, width=Inches(6.0))
        except Exception as e:
            p = doc.add_paragraph(f"【附件转换失败: {str(e)}】")
            p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0)

    # 保存到内存字节流
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    
    return byte_io
